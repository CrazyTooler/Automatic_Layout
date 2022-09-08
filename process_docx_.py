# 处理word文档
from time import time
import docx
import os
import re
from PIL import Image
from functools import reduce
import math

class Article:
    """构造Article类
    属性：title，paragraphs，count_pics，
          have_pics,count_subtitle"""
    def __init__(self,title,paras=[],count_pics=0):
        self.title = title
        self.sub_title = []
        self.paragraphs = paras
        self.count_pics = count_pics
        self.have_subtitle = False
        self.count_sub_title = 0
        self.count_title_char = 0
        self.count_paras_char = 0
        self.count_subtitle_char = 0
        self.count_paras = 0
        self.article_spaces = []
        self.pic =[]
        self.ishead = False
        if count_pics > 0:
            self.have_pics = True
        else:
            self.have_pics = False

    def get_picture(self,pic_size):
        if self.have_pics:
            self.pic = pic_size[0:self.count_pics]
            pic_size[0:self.count_pics] = []

    def set_count(self):
        """设置标题和正文字数"""
        self.para_char=[]
        self.count_title_char = len(self.title)
        self.count_ftitle_char = len(self.sub_title[0])
        self.count_ytitle_char = len(self.sub_title[1])
        if self.have_subtitle:
            self.count_subtitle_char = len(reduce(lambda x,y: x+y,self.sub_title))
        else:
            self.count_subtitle_char = 0
        count = 0
        
        for para in self.paragraphs:
            numcount=0
            for i in range(10):
                numcount+=para.count(str(i))
            count += len(para)
            self.para_char.append(len(para)-numcount//2-1)
        self.count_paras_char = count-numcount//2-1

    def set_default_size(self):
        """选取默认样式，计算默认大小"""
        pt = 0.35
        line_hight = 1.3*10*pt
        char_width = 10*pt 
        width_series = [29.9,40,70,100]
        for width in width_series:
            width_ = width - 10/3.2#文本块与边框之间的距离为5mm
            pic_space = 0
            for pic in self.pic:
                pic_space += 1600*pic[1] // pic[0]#默认图片宽度为40
            self.pics_space = pic_space#图片
            maintitle_rows = math.ceil(len(self.title)*40*pt/(width_*3.2))
            subtitle_rows = math.ceil(len(self.sub_title[0])*20*pt/(width_*3.2)) + math.ceil(len(self.sub_title[1])*20*pt/(width_*3.2))
            title_hight = maintitle_rows*40*pt + subtitle_rows*20*pt#标题高度
            text_rows = 0
            column_count = math.floor((width+10)/20)#栏数
            for para in self.paragraphs:
                text_rows += math.ceil((len(para)+2)*char_width/17/3.2  )
            text_rows = text_rows // column_count
            # print(f'ddd----{self.title}-------------')
            # print(maintitle_rows)
            # print(subtitle_rows)
            # print(text_rows)
            text_hight = text_rows*line_hight
            self.title_space = title_hight * width // 5
            self.text_space = width * text_hight // 5 
            self.article_space = self.title_space + self.text_space + self.pics_space
            self.article_spaces.append(self.article_space)

def get_title_and_para(docx_path):
    """获取word文本标题以及段落内容
    返回字典（index：article）"""
    result_dict = dict()
    count = 0
    file = docx.Document(docx_path)
    for para in file.paragraphs:#遍历每一段内容，为title和paragraphs，以及图片注释
        if para.text:
            if para.style.name.startswith('Heading 1'):#用1级标题表示新闻大标题
                count += 1
                atcl = Article(para.text,[])
                result_dict[count] = atcl
            else:
                if para.style.name.startswith('Heading 2'):#用2级标题表示副标题
                    atcl.sub_title.append(para.text)
                    atcl.have_subtitle = True
                    atcl.count_sub_title += 1
                elif para.text.startswith(f'图片{count}'):
                    atcl.have_pics = True
                    atcl.count_pics += 1
                else:
                    atcl.paragraphs.append(para.text)
    return result_dict

def get_pictures(word_path, result_path,imglist):
    """从word中获得图片，存储在目标路径
    图片命名为文件名_图片名（image+n）"""
    doc = docx.Document(word_path)
    dict_rel = doc.part._rels
    if not os.path.exists(result_path):
        os.makedirs(result_path)
    else:
        for file in os.listdir(result_path):
            os.remove(result_path+'/'+file)
    for rel in dict_rel:
        rel = dict_rel[rel]
        if "image" in rel.target_ref:#检查word是否含有图片
            img_name = re.findall("/(.*)", rel.target_ref)[0]
            word_name = os.path.splitext(word_path)[0]
            if os.sep in word_name:
                new_name = word_name.split('/')[-1]
            else:
                new_name = word_name.split('/')[-1]
            img_name = f'{new_name}_{img_name}'#命名图片
            with open(f'{result_path}/{img_name}', "wb") as f:
                f.write(rel.target_part.blob)
    img_paths = ['./img_article/'+name for name in os.listdir('./img_article')]
    result = []
    for img_path in img_paths:
        img = Image.open(img_path)
        imgSize = img.size  #大小/尺寸
        w = img.width       #图片的宽
        h = img.height      #图片的高
        result.append((w,h,img_path))
        img.close()#关闭图片
        '''
    for img_path in img_paths:#重命名
        suffix = imglist.pop(0)
        new = img_path.split('.')[-1]
        os.rename(img_path,'./img_article/'+new_name+suffix+'.'+new)
        '''
    return result

def prn_obj(obj): 
    """打印对象属性"""
    print ('\n'.join(['%s:%s' % item for item in obj.__dict__.items()]) )

def get_article_dic(filename:str):
    result = get_title_and_para(filename)
    ########
    a = [0]
    for i in range(len(result)):
        a.append(result[i+1].count_pics)
    s = list(a)
    imglist = []
    for i in range(len(s)):
        j = 1
        while s[i] > 0:
            imglist.append(f'_{i}_{j}_')
            s[i] -= 1
            j += 1
    ########
    pic_size = get_pictures(filename,'./img_article',imglist)
    for n in result:
        result[n].set_count()
        result[n].get_picture(pic_size)
        ###
        # prn_obj(result[n])
        ###
        for i in range(2):
            if 'None' in result[n].sub_title[i]:
                result[n].sub_title[i] = ''
                result[n].count_sub_title -= 1
        if result[n].sub_title[0]=='' and result[n].sub_title[1]=='':
            result[n].have_subtitle=False
        # result[n].set_default_size()
        # print('frfrfrfrffffffffffffffffff',result[n].para)
        result[n].count_paras = len(result[n].paragraphs)
        result[n].count_ftitle_char = len(result[n].sub_title[0])
        result[n].count_ytitle_char = len(result[n].sub_title[1])
    return result

if __name__ == '__main__':
    #E:\vscodeproject\test1\python_tex\template_initiate\testdocx\test.docx
    sss = time()
    result = get_title_and_para('./testdocx/jh2.docx')
    ##########
    a = [0]
    for i in range(len(result)):
        a.append(result[i+1].count_pics)
    s = list(a)
    imglist = []
    for i in range(len(s)):
        j = 1
        while s[i] > 0:
            imglist.append(f'_{i}_{j}_')
            s[i] -= 1
            j += 1
    ##########
    pic_size = get_pictures('./testdocx/jh2.docx','./img_article',imglist)
    print(pic_size)
    print(len(result))
    for n in result:
        result[n].set_count()
        result[n].get_picture(pic_size)
        for i in range(2):
            if 'None' in result[n].sub_title[i]:
                result[n].sub_title[i] = ''
                result[n].count_sub_title -= 1
        if result[n].sub_title[0]=='' and result[n].sub_title[1]=='':
            result[n].have_subtitle=False
        # result[n].set_default_size()
        result[n].count_paras = len(result[n].paragraphs)
        result[n].count_ftitle_char = len(result[n].sub_title[0])
        result[n].count_ytitle_char = len(result[n].sub_title[1])
        prn_obj(result[n])
        print('----------------------------------')
    print(time()-sss)
